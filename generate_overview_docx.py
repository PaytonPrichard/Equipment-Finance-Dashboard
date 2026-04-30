"""Generate Tranche Platform Overview as a formatted Word document."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

GOLD = RGBColor(0xC8, 0xA5, 0x4A)
DARK = RGBColor(0x11, 0x18, 0x27)
GRAY = RGBColor(0x6B, 0x72, 0x80)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

def set_cell_shading(cell, color_hex):
    from docx.oxml.ns import qn
    from lxml import etree
    shading = etree.SubElement(cell._tc.get_or_add_tcPr(), qn('w:shd'))
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')

def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = DARK
    return h

def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.style.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.color.rgb = GRAY
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = DARK
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = GRAY
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = GRAY
    return p

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ── Title ──
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = title.add_run('TRANCHE')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = DARK
title.add_run('\n')
run2 = title.add_run('ABL Deal Screening Platform')
run2.font.size = Pt(14)
run2.font.color.rgb = GOLD

sub = doc.add_paragraph()
run = sub.add_run('One-page overview for credit teams evaluating the platform.')
run.font.size = Pt(11)
run.font.color.rgb = GRAY
run.font.italic = True

doc.add_paragraph()  # spacer

# ── What is Tranche? ──
add_heading_styled(doc, 'What is Tranche?', level=2)
add_body(doc, 'Tranche is a deal screening tool for asset-based lenders. It takes raw deal inputs (borrower financials, collateral details, deal structure) and produces a scored risk assessment with a pass/flag/fail verdict in under 2 minutes.')
add_body(doc, 'Built for pre-origination screening, not underwriting. The goal is to tell your team which deals are worth pursuing before committing analyst time.')

# ── Inputs ──
add_heading_styled(doc, 'What inputs does it take?', level=2)

add_body(doc, 'Borrower profile (shared across all asset classes):')
add_body(doc, 'Company name, years in business, annual revenue, EBITDA, total existing debt, industry sector, credit rating.')

p = doc.add_paragraph()
run = p.add_run('Equipment Finance: ')
run.bold = True
run.font.size = Pt(11)
run.font.color.rgb = DARK
run = p.add_run('Equipment type, condition, cost, down payment, financing type (EFA/FMV/TRAC), useful life, loan term, essential-use flag.')
run.font.size = Pt(11)
run.font.color.rgb = GRAY

p = doc.add_paragraph()
run = p.add_run('Accounts Receivable: ')
run.bold = True
run.font.size = Pt(11)
run.font.color.rgb = DARK
run = p.add_run('Total AR outstanding, aging buckets (current/30+/60+/90+), top customer concentration, dilution rate, ineligibles %, advance rate.')
run.font.size = Pt(11)
run.font.color.rgb = GRAY

p = doc.add_paragraph()
run = p.add_run('Inventory Finance: ')
run.bold = True
run.font.size = Pt(11)
run.font.color.rgb = DARK
run = p.add_run('Total inventory, breakdown (raw/WIP/finished/obsolete), turnover, days on hand, NOLV %, advance rate, perishable flag.')
run.font.size = Pt(11)
run.font.color.rgb = GRAY

# ── Outputs ──
add_heading_styled(doc, 'What does it output?', level=2)

add_bullet(doc, ' Composite weighted score across 7 factors per asset class.', bold_prefix='Risk score (0-100).')
add_bullet(doc, ' Pass (75+), Flag (35-74), or Fail (<35) with specific reasons.', bold_prefix='Verdict.')
add_bullet(doc, ' DSCR, leverage, screening rate, debt service, LTV, borrowing base, DSO, concentration.', bold_prefix='Key metrics.')
add_bullet(doc, ' 4-scenario EBITDA stress (base, -10%, -20%, -30%) showing score/DSCR degradation.', bold_prefix='Stress test.')
add_bullet(doc, ' Suggested deal terms and structure based on the risk profile.', bold_prefix='Structure recommendation.')
add_bullet(doc, ' One-click PDF with firm logo, colors, and disclaimers. Committee-ready.', bold_prefix='Branded memo.')

# ── Scoring Logic ──
add_heading_styled(doc, 'What is the screening logic based on?', level=2)
add_body(doc, 'Rules-based scoring model. No AI, no black box. Each asset class scores 7 weighted factors.')

# Scoring table
table = doc.add_table(rows=8, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.LEFT
table.style = 'Table Grid'

headers = ['Factor', 'Weight', 'What it measures']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = WHITE
    set_cell_shading(cell, '111827')

rows_data = [
    ('DSCR', '25%', 'Debt service coverage. Higher is better.'),
    ('Leverage', '20%', 'Total debt / EBITDA. Lower is better.'),
    ('Industry Risk', '15%', 'Sector-based risk tier (low/moderate/high).'),
    ('Essential Use', '10%', 'Whether equipment is critical to operations.'),
    ('Equipment LTV', '10%', 'Loan-to-value with condition adjustment.'),
    ('Years in Business', '10%', 'Operating track record.'),
    ('Term Coverage', '10%', 'Loan term vs. useful life ratio.'),
]
for r, (factor, weight, desc) in enumerate(rows_data, start=1):
    for c, val in enumerate([factor, weight, desc]):
        cell = table.rows[r].cells[c]
        cell.text = val
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = GRAY
    if r % 2 == 0:
        for c in range(3):
            set_cell_shading(table.rows[r].cells[c], 'F9FAFB')

doc.add_paragraph()
add_body(doc, 'AR and Inventory models swap in asset-specific factors (aging, concentration, dilution for AR; quality, composition, NOLV for inventory) while keeping DSCR, leverage, and experience as shared factors.')
add_body(doc, 'Pass/flag/fail thresholds are configurable per firm. Defaults: 1.25x min DSCR, 5.0x max leverage, 75+ to pass.')

# ── AI? ──
add_heading_styled(doc, 'Any AI component?', level=2)
add_body(doc, 'No. The model is entirely deterministic and rules-based. Same inputs always produce the same score. Every factor weight and threshold is visible and configurable.')
add_body(doc, 'This is intentional. Credit committees want to understand exactly why a deal scored the way it did. A transparent, auditable model builds trust faster than a black-box prediction.')

# ── Time Savings ──
add_heading_styled(doc, 'What saves an analyst the most time?', level=2)

add_bullet(doc, ' Upload a CSV of 100+ deals. Every deal scored and ranked in seconds.', bold_prefix='Batch screening.')
add_bullet(doc, ' Enter data, get a verdict. No building models from scratch.', bold_prefix='Instant pass/flag/fail.')
add_bullet(doc, ' One click generates a committee-ready memo. No copy-pasting into Word.', bold_prefix='Branded PDF memos.')
add_bullet(doc, ' Deals move through stages. Stage changes trigger email notifications.', bold_prefix='Pipeline tracking.')
add_bullet(doc, ' REST API and webhooks connect to any CRM. No double entry.', bold_prefix='CRM integration.')

# ── Onboarding ──
add_heading_styled(doc, 'How fast is onboarding?', level=2)
add_body(doc, 'Upload your pipeline CSV, set your credit policy thresholds, invite your team. Dashboard is populated in one afternoon. No implementation project. No consultants.')

# ── Footer ──
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('Tranche')
run.font.bold = True
run.font.size = Pt(11)
run.font.color.rgb = DARK
run = footer.add_run('  |  gettranche.app  |  team@gettranche.app')
run.font.size = Pt(11)
run.font.color.rgb = GRAY

# Save
output_path = 'outputs/Tranche_Platform_Overview.docx'
doc.save(output_path)
print(f'Saved to {output_path}')
