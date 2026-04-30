"""Convert Deal_Screening_Model_Assumptions.md to a formatted Word document."""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def parse_table(lines):
    """Parse markdown table lines into rows of cells."""
    rows = []
    for line in lines:
        line = line.strip()
        if not line.startswith("|"):
            continue
        cells = [c.strip() for c in line.split("|")[1:-1]]
        # Skip separator rows (---|---)
        if all(re.match(r'^[-:]+$', c) for c in cells):
            continue
        rows.append(cells)
    return rows


def add_table_to_doc(doc, rows):
    """Add a formatted table to the document."""
    if not rows:
        return
    n_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=n_cols, style='Light Grid Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            cell = table.cell(i, j)
            # Strip bold markers for display
            clean = cell_text.replace("**", "")
            cell.text = clean
            for paragraph in cell.paragraphs:
                paragraph.style = doc.styles['Normal']
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    if i == 0:  # Header row
                        run.bold = True

    doc.add_paragraph()  # spacing after table


def process_inline(paragraph, text):
    """Handle bold (**), italic (*), and inline code (`) in text."""
    # Pattern: chunks of **bold**, *italic*, `code`, or plain text
    pattern = r'(\*\*.*?\*\*|\*.*?\*|`.*?`|[^*`]+)'
    parts = re.findall(pattern, text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)
        else:
            paragraph.add_run(part)


def convert(md_path, docx_path):
    md_text = Path(md_path).read_text(encoding="utf-8")
    lines = md_text.splitlines()

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip horizontal rules
        if stripped == '---':
            i += 1
            continue

        # Code blocks
        if stripped.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            p = doc.add_paragraph()
            run = p.add_run("\n".join(code_lines))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            p.paragraph_format.left_indent = Inches(0.5)
            continue

        # Tables: collect consecutive lines starting with |
        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = parse_table(table_lines)
            add_table_to_doc(doc, rows)
            continue

        # Headings
        if stripped.startswith("#"):
            match = re.match(r'^(#{1,6})\s+(.*)', stripped)
            if match:
                level = len(match.group(1))
                heading_text = match.group(2).replace("**", "")
                if level == 1:
                    doc.add_heading(heading_text, level=0)
                else:
                    doc.add_heading(heading_text, level=min(level, 4))
                i += 1
                continue

        # Bullet / list items
        if re.match(r'^[-*]\s', stripped):
            text = re.sub(r'^[-*]\s+', '', stripped)
            p = doc.add_paragraph(style='List Bullet')
            process_inline(p, text)
            i += 1
            continue

        # Numbered list
        if re.match(r'^\d+\.\s', stripped):
            text = re.sub(r'^\d+\.\s+', '', stripped)
            p = doc.add_paragraph(style='List Number')
            process_inline(p, text)
            i += 1
            continue

        # Empty lines
        if not stripped:
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        process_inline(p, stripped)
        i += 1

    doc.save(docx_path)
    print(f"Saved: {docx_path}")


if __name__ == "__main__":
    base = Path(__file__).parent
    convert(
        base / "Deal_Screening_Model_Assumptions.md",
        base / "outputs" / "Deal_Screening_Model_Assumptions.docx",
    )
